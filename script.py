import os
import re
import argparse
from docx import Document
from PIL import Image
from transformers import AutoModel, AutoTokenizer
import fitz  # PyMuPDF
import tempfile
import warnings
warnings.filterwarnings("ignore")



#Using GOT-OCR model
#https://huggingface.co/stepfun-ai/GOT-OCR2_0

tokenizer = AutoTokenizer.from_pretrained('ucaslcl/GOT-OCR2_0', trust_remote_code=True)
model = AutoModel.from_pretrained(
    'ucaslcl/GOT-OCR2_0', 
    trust_remote_code=True, 
    low_cpu_mem_usage=True, 
    device_map='cuda', #NVidia GPU must
    use_safetensors=True, 
    pad_token_id=tokenizer.eos_token_id
    )
model = model.eval().cuda()
print('Model loading complete!\n')

def extract_text(pdf_path, docx_path, latex=True, page_break=True):
    """
    Converting each page of the PDF file to text and save the result as .docx file
    """
    pdf_path = os.path.abspath(pdf_path) #get absolute path
    print(f"PDF path: {pdf_path}")
    
    if not os.path.exists(pdf_path):
        print(f"Error: PDF file not found at {pdf_path}")
        return
    try:
        pdf = fitz.open(pdf_path) # open the file
    except Exception as e:
        print(f"Error opening PDF: {e}")
        return
    
    doc = Document() #create an empty document
    
    for page_num in range(len(pdf)):
        page = pdf[page_num]
        pix = page.get_pixmap() #convert the page to image
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples) #convert to PIL image

        #save current image as a temporary file
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_file:
            img.save(temp_file.name)
            temp_file_path = temp_file.name

        # OCR using GOT-OCR model
        try: 
            res = model.chat(tokenizer, temp_file_path, ocr_type='format')
            doc.add_paragraph(res)
            if latex:
                doc.add_paragraph("\\newpage") #add \newpage after every page for LaTeX
            if page_break:
                # Add a page break after each image result, except for the last one
                if page_num < len(pdf) - 1:
                    doc.add_page_break()
                    
            print(f"OCR of page {page_num} done.")
        
        except Exception as e:
            print(f"Error performing OCR on page {page_num}: {e}")

    
    pdf.close() #close the PDF 
    doc.save(docx_path)
    print(f"Text Output saved to {docx_path}")

def docx_to_latex(docx_path, latex_output_path):
    """
    Convert a .docx file to a LaTeX .tex file.
    :param docx_path: Path to the input .docx file.
    :param latex_output_path: Path to save the output .tex file.
    """

    try:
        doc = Document(docx_path) # Load the .docx file

        # LaTeX preamble
        latex_content = """\\documentclass[14pt]{extarticle}
\\usepackage[utf8]{inputenc}
\\usepackage{amsmath}
\\usepackage{amssymb}
\\usepackage{geometry}
\\usepackage{fontspec}
\\setmainfont{Times New Roman}
\\geometry{
a4paper,
left=1in,
right=1in,
top=0.6in,
bottom=0.6in
}
\\begin{document}
"""

        #concatenate all paragraphs into a single string
        full_text = "\n".join(paragraph.text.strip() for paragraph in doc.paragraphs)

        #replace \title{} blocks with \section*{} blocks and remove newlines inside the title block
        full_text = re.sub(r'\\title\{([^}]*)\}', lambda m: '\\section*{'+ m.group(1).replace('\n', ' ') + '}', full_text)

        #replace \author{} blocks and remove newlines and extra backslashes inside the author block
        full_text = re.sub(r'\\author\{([^}]*)\}', lambda m: '\\textbf{' + m.group(1).replace('\\\\', '').replace('\n', ' ').strip() + '}', full_text)

        # remove newlines and extra backslashes inside the footnotetext block
        full_text = re.sub(r'\\footnotetext\{([^}]*)\}', lambda m: '\\footnotetext{'+ m.group(1).replace('\n', '') + '}', full_text)

        # replace \n with \\\n\n
        latex_content += full_text.replace('\n', '\\\\\n\n') 
        # replace \newpage\\\n\n\\ with \newpage
        latex_content = latex_content.replace("\\newpage\\\\\n\n\\\\", "\\newpage") 
        # replace \\\n\n\\ with \\
        latex_content = latex_content.replace("\\\\\n\n\\\\\n", "\\\\\n" ) 
        # replace \\\n\n\newpage\n\ with \n\n\newpage\n
        latex_content = latex_content.replace("\\\\\n\n\\newpage\n", "\n\n\\newpage\n") 
        #replace }\\ with }
        latex_content = latex_content.replace("}\\\\", "}" ) 
        #remove the marker from footnotes
        latex_content = latex_content.replace("\\footnotetext", "\let\\thefootnote\\relax\\footnotetext")
        
        latex_content += "\n\n\\end{document}\n"

        # save the LaTeX content to a .tex file
        with open(latex_output_path, "w", encoding="utf-8") as tex_file:
            tex_file.write(latex_content)
        
        print(f'LaTeX output save to {latex_output_path}')

    except Exception as e:
        print(f"An error occurred: {str(e)}")

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Convert PDF to DOCX and LaTeX')
    parser.add_argument('--pdf_name', type=str, required=True, help='Name of the PDF file')
    parser.add_argument('--doc_name', type=str, default='doc_output.docx', help='Name of the output .docx file (default: doc_output.docx)')
    parser.add_argument('--tex_name', type=str, default='tex_output.tex', help='Name of the output .tex file (default: tex_output.tex)')
    args = parser.parse_args()

    pdf_name = args.pdf_name 
    doc_name = args.doc_name 
    tex_name = args.tex_name 
    
    print(f'Converting {pdf_name} to DOCX and LaTeX...')
    print(f'DOCX output will be saved to {doc_name}')
    print(f'LaTeX output will be saved to {tex_name}')

    current_dir = os.getcwd() # Get the current working directory
    pdf_full_path = os.path.join(current_dir, pdf_name) # Create full path

    extract_text(pdf_full_path, doc_name, latex=False, page_break=False)
    docx_to_latex(doc_name, tex_name)